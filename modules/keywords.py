import pandas as pd
import string
import spacy
import docx  # pip install python-docx
import PyPDF2
from spacy.lang.en.stop_words import STOP_WORDS
from spacy.lang.en import English
import hashlib

nlp = spacy.load("en_core_web_sm")
parser = English()
punctuations = string.punctuation
stop_words = set(STOP_WORDS)


def text_cleaning(docx):
    mytokens = [word.lemma_.lower().strip() if word.lemma_ != "-PRON-" else word.lower_ for word in docx]
    mytokens = [word for word in mytokens if word not in stop_words and word not in punctuations]
    mytokens = [word for word in mytokens if word.isalpha()]
    mytokens = " ".join(mytokens)
    return mytokens


def compute_term_frequency(word_dictionary, bag_of_words):
    term_frequency_dictionary = {}
    length_of_bag_of_words = len(bag_of_words)

    for word, count in word_dictionary.items():
        term_frequency_dictionary[word] = count / float(length_of_bag_of_words)

    return term_frequency_dictionary


# def create_hashing(key_list):
#     keywords_list1 = []
#     for i in key_list:
#         h1 = hashlib.sha256()
#         h1.update(i.encode('utf-8'))
#         txt1 = h1.hexdigest()
#         keywords_list1.append(txt1)
#     return keywords_list1


def generate_keywords(key_file):
    if key_file.endswith('.txt'):
        with open(file=key_file, mode='r') as file:
            text = file.read()

    elif key_file.endswith('.docx'):
        doc = docx.Document(key_file)
        text = ''
        for i in range(len(doc.paragraphs)):
            text += doc.paragraphs[i].text

    elif key_file.endswith('.pdf'):
        with open(key_file, 'rb') as file:
            pdfReader = PyPDF2.PdfFileReader(file)
            pageObj = pdfReader.getPage(0)
            text = pageObj.extractText()

    list_sents = []
    doc = nlp(text)
    for sent in doc.sents:
        list_sents.append(sent)
    sentences = []
    for sente in list_sents:
        sentences.append(text_cleaning(sente))

    word_count = {}
    for i in range(len(sentences)):
        for word in sentences[i].split():
            if word in word_count:
                word_count[word] += 1
            else:
                word_count[word] = 1

    thresh = 3
    vocab = {}
    for word, count in word_count.items():
        if count >= thresh:
            vocab[word] = count

    df_dic = compute_term_frequency(vocab, vocab.keys())
    df = pd.DataFrame()
    df['Tokens'] = df_dic.keys()
    df['WordImportance'] = df_dic.values()
    df = df.sort_values(by="WordImportance", ascending=False).reset_index(drop=True)
    # df.to_csv("WordWithImportances.csv", index=False)

    keywords_list = list(df['Tokens'].values)
    keywords_weight = list(df['WordImportance'].values)

    # keywords_list1 = create_hashing(keywords_list)
    return keywords_list, keywords_weight
